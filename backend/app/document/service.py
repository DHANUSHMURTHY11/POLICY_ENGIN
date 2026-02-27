"""
Document Composer service — Word, PDF, JSON generation via AI composition.
Strict AI Document Mode — no fallback, no generic enhancer.
AI composes narratives from raw structure before rendering.
"""
import json
import os
import uuid
from datetime import datetime, timezone
from typing import Any, List, Optional

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.core.logging import get_logger
from app.database.mongodb import policy_documents_collection
from app.ai.providers import AIProviderError
from app.document.ai_document_composer import ai_document_composer
from app.document.schemas import AIComposedDocument, ApprovalFlowEntry
from app.policy.models import PolicyMetadata

logger = get_logger(__name__)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "generated_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════════

async def _get_latest_structure(policy_id: str) -> dict:
    """Fetch the latest document_structure from MongoDB."""
    collection = policy_documents_collection()
    doc = await collection.find_one(
        {"policy_id": policy_id},
        sort=[("version", -1)],
    )
    if not doc or "document_structure" not in doc:
        raise ValueError(f"No document structure found for policy {policy_id}")
    return doc["document_structure"]


async def _check_policy_approved(db: AsyncSession, policy_id: str) -> PolicyMetadata:
    """Verify that the policy exists and is approved.
    Raises 404 if not found, 403 if not approved.
    """
    result = await db.execute(
        select(PolicyMetadata).where(PolicyMetadata.id == uuid.UUID(policy_id))
    )
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(status_code=404, detail="Policy not found")
    if row.status != "approved":
        raise HTTPException(
            status_code=403,
            detail=f"Document generation blocked: policy status is '{row.status}'. Only approved policies can generate documents.",
        )
    return row


async def _get_approval_flow(db: AsyncSession, policy_id: str) -> List[dict]:
    """Fetch the approval workflow data for a policy."""
    try:
        from app.workflow.models import PolicyWorkflowInstance
        result = await db.execute(
            select(PolicyWorkflowInstance)
            .options(selectinload(PolicyWorkflowInstance.actions))
            .where(PolicyWorkflowInstance.policy_id == uuid.UUID(policy_id))
            .order_by(PolicyWorkflowInstance.created_at.desc())
            .limit(1)
        )
        instance = result.scalar_one_or_none()
        if not instance:
            return []

        flow = []
        for action in sorted(instance.actions, key=lambda a: a.created_at):
            flow.append({
                "level": action.level_number if hasattr(action, "level_number") else 0,
                "role": action.role if hasattr(action, "role") else "Approver",
                "approver": str(action.user_id) if hasattr(action, "user_id") else "",
                "status": action.action if hasattr(action, "action") else "approved",
                "timestamp": str(action.created_at) if hasattr(action, "created_at") else "",
                "comments": action.comments if hasattr(action, "comments") else "",
            })
        return flow
    except Exception as exc:
        logger.warning(
            f"Could not fetch approval flow: {exc}",
            extra={"event": "approval_flow_fetch_warning", "policy_id": policy_id},
        )
        return []


async def _compose_via_ai(
    structure: dict, approval_flow: List[dict], policy_id: str
) -> AIComposedDocument:
    """Send structure to AI for document composition. Raises 500 on failure."""
    try:
        return await ai_document_composer.compose_document(structure, approval_flow)
    except AIProviderError as exc:
        logger.error(
            "AI document composition failed",
            extra={
                "event": "ai_compose_error",
                "policy_id": policy_id,
                "error": str(exc),
            },
        )
        raise HTTPException(
            status_code=500,
            detail="AI document composition failed. Cannot generate document without AI.",
        )


# ═══════════════════════════════════════════════════════════════════
#  Word (.docx) Generation — AI-Composed
# ═══════════════════════════════════════════════════════════════════

async def generate_word(
    db: AsyncSession, policy_id: str, filename_prefix: str
) -> str:
    """Generate a Word document from AI-composed policy content.
    Step 1: Check approval → Step 2: AI compose → Step 3: Render Word.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Step 1: Approval gate
    await _check_policy_approved(db, policy_id)

    # Step 2: Fetch structure + approval flow → AI compose
    structure = await _get_latest_structure(policy_id)
    approval_flow = await _get_approval_flow(db, policy_id)
    composed = await _compose_via_ai(structure, approval_flow, policy_id)

    # Step 3: Render Word from composed document
    doc = DocxDocument()

    # ── Title Page ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run(composed.title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    header = structure.get("header", {})
    if header.get("organization"):
        org_para = doc.add_paragraph()
        org_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        org_run = org_para.add_run(header["organization"])
        org_run.font.size = Pt(14)

    if header.get("effective_date"):
        doc.add_paragraph(f"Effective Date: {header['effective_date']}")
    if header.get("expiry_date"):
        doc.add_paragraph(f"Expiry Date: {header['expiry_date']}")

    doc.add_page_break()

    # ── Table of Contents ──
    doc.add_heading("Table of Contents", level=1)
    for i, section in enumerate(composed.sections, 1):
        toc_para = doc.add_paragraph(f"{i}. {section.heading}")
        toc_para.paragraph_format.left_indent = Inches(0.5)
    doc.add_page_break()

    # ── Scope ──
    if composed.scope:
        doc.add_heading("Scope", level=1)
        doc.add_paragraph(composed.scope)
        doc.add_paragraph("")

    # ── Sections with narratives ──
    for i, section in enumerate(composed.sections, 1):
        doc.add_heading(f"{i}. {section.heading}", level=1)
        if section.content:
            doc.add_paragraph(section.content)

        # Render tables
        for tbl in section.tables:
            if tbl.get("caption"):
                cap_para = doc.add_paragraph()
                cap_run = cap_para.add_run(tbl["caption"])
                cap_run.bold = True
                cap_run.font.size = Pt(10)

            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                for idx, h in enumerate(headers):
                    table.rows[0].cells[idx].text = str(h)
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for idx, val in enumerate(row_data):
                        if idx < len(row_cells):
                            row_cells[idx].text = str(val)

        doc.add_paragraph("")  # spacing

    # ── Approval Flow ──
    if composed.approval_flow_summary or composed.approval_chain:
        doc.add_page_break()
        doc.add_heading("Approval Flow", level=1)

        if composed.approval_flow_summary:
            doc.add_paragraph(composed.approval_flow_summary)
            doc.add_paragraph("")

        if composed.approval_chain:
            atbl = doc.add_table(rows=1, cols=5)
            atbl.style = "Light Grid Accent 1"
            ah = atbl.rows[0].cells
            ah[0].text = "Level"
            ah[1].text = "Role"
            ah[2].text = "Status"
            ah[3].text = "Timestamp"
            ah[4].text = "Comments"
            for entry in composed.approval_chain:
                arow = atbl.add_row().cells
                arow[0].text = str(entry.level)
                arow[1].text = entry.role
                arow[2].text = entry.status
                arow[3].text = entry.timestamp or ""
                arow[4].text = entry.comments

    # ── Annexures ──
    if composed.annexures:
        doc.add_page_break()
        doc.add_heading("Annexures", level=1)
        for annex in composed.annexures:
            if isinstance(annex, dict):
                doc.add_heading(annex.get("title", "Annexure"), level=2)
                doc.add_paragraph(annex.get("content", ""))

    # ── Version History ──
    version_control = structure.get("version_control", [])
    if version_control:
        doc.add_page_break()
        doc.add_heading("Version History", level=1)
        vtable = doc.add_table(rows=1, cols=4)
        vtable.style = "Light Grid Accent 1"
        vh = vtable.rows[0].cells
        vh[0].text = "Version"
        vh[1].text = "Date"
        vh[2].text = "Author"
        vh[3].text = "Change Summary"
        for vc in version_control:
            vrow = vtable.add_row().cells
            vrow[0].text = str(vc.get("version_number", ""))
            vrow[1].text = str(vc.get("created_at", ""))
            vrow[2].text = vc.get("created_by", "")
            vrow[3].text = vc.get("change_summary", "")

    filepath = os.path.join(OUTPUT_DIR, f"{filename_prefix}_{uuid.uuid4().hex[:8]}.docx")
    doc.save(filepath)

    logger.info(
        "Word document generated (AI-composed)",
        extra={
            "event": "document_generated",
            "format": "docx",
            "policy_id": policy_id,
            "ai_composed": True,
        },
    )
    return filepath


# ═══════════════════════════════════════════════════════════════════
#  PDF Generation — AI-Composed
# ═══════════════════════════════════════════════════════════════════

async def generate_pdf(
    db: AsyncSession, policy_id: str, filename_prefix: str
) -> str:
    """Generate a PDF from AI-composed policy content.
    Step 1: Check approval → Step 2: AI compose → Step 3: Render PDF.
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    # Step 1: Approval gate
    await _check_policy_approved(db, policy_id)

    # Step 2: Fetch structure + approval flow → AI compose
    structure = await _get_latest_structure(policy_id)
    approval_flow = await _get_approval_flow(db, policy_id)
    composed = await _compose_via_ai(structure, approval_flow, policy_id)

    filepath = os.path.join(OUTPUT_DIR, f"{filename_prefix}_{uuid.uuid4().hex[:8]}.pdf")
    pdf = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    # Custom styles
    title_style = ParagraphStyle(
        "PolicyTitle", parent=styles["Title"],
        fontSize=24, spaceAfter=12, textColor=colors.HexColor("#003366"),
    )
    h1_style = ParagraphStyle(
        "PolicyH1", parent=styles["Heading1"],
        fontSize=16, spaceAfter=8, textColor=colors.HexColor("#003366"),
    )
    h2_style = ParagraphStyle(
        "PolicyH2", parent=styles["Heading2"],
        fontSize=13, spaceAfter=6, textColor=colors.HexColor("#1a5276"),
    )
    body_style = ParagraphStyle(
        "PolicyBody", parent=styles["Normal"],
        fontSize=10, spaceAfter=8, leading=14,
    )

    # ── Title ──
    header = structure.get("header", {})
    elements.append(Paragraph(composed.title, title_style))
    if header.get("organization"):
        elements.append(Paragraph(header["organization"], styles["Normal"]))
    if header.get("effective_date"):
        elements.append(Paragraph(f"Effective Date: {header['effective_date']}", styles["Normal"]))
    elements.append(Spacer(1, 24))

    # ── TOC ──
    elements.append(Paragraph("Table of Contents", h1_style))
    for i, section in enumerate(composed.sections, 1):
        elements.append(Paragraph(f"{i}. {section.heading}", styles["Normal"]))
    elements.append(PageBreak())

    # ── Scope ──
    if composed.scope:
        elements.append(Paragraph("Scope", h1_style))
        elements.append(Paragraph(composed.scope, body_style))
        elements.append(Spacer(1, 12))

    # ── Sections with narratives ──
    for i, section in enumerate(composed.sections, 1):
        elements.append(Paragraph(f"{i}. {section.heading}", h1_style))
        if section.content:
            elements.append(Paragraph(section.content, body_style))
        elements.append(Spacer(1, 8))

        # Render tables
        for tbl in section.tables:
            if tbl.get("caption"):
                elements.append(Paragraph(tbl["caption"], h2_style))

            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers:
                col_count = len(headers)
                col_width = 6.5 * inch / col_count
                table_data = [headers] + rows
                t = Table(table_data, colWidths=[col_width] * col_count)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ]))
                elements.append(t)
                elements.append(Spacer(1, 12))

    # ── Approval Flow ──
    if composed.approval_flow_summary or composed.approval_chain:
        elements.append(PageBreak())
        elements.append(Paragraph("Approval Flow", h1_style))

        if composed.approval_flow_summary:
            elements.append(Paragraph(composed.approval_flow_summary, body_style))
            elements.append(Spacer(1, 8))

        if composed.approval_chain:
            af_data = [["Level", "Role", "Status", "Timestamp", "Comments"]]
            for entry in composed.approval_chain:
                af_data.append([
                    str(entry.level),
                    entry.role,
                    entry.status,
                    entry.timestamp or "",
                    entry.comments,
                ])
            af_table = Table(af_data, colWidths=[0.6 * inch, 1.2 * inch, 0.8 * inch, 1.5 * inch, 2.4 * inch])
            af_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a5276")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
            ]))
            elements.append(af_table)

    # ── Annexures ──
    if composed.annexures:
        elements.append(PageBreak())
        elements.append(Paragraph("Annexures", h1_style))
        for annex in composed.annexures:
            if isinstance(annex, dict):
                elements.append(Paragraph(annex.get("title", "Annexure"), h2_style))
                elements.append(Paragraph(annex.get("content", ""), body_style))

    # Page drawing callback for Footers
    version_control = structure.get("version_control", [])
    current_version = version_control[-1].get("version_number", 1) if version_control else 1

    def draw_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        footer_text = f"Official Confidential | Version: {current_version} | Page {doc.page}"
        # Draw at bottom center
        canvas.drawCentredString(A4[0] / 2.0, 0.5 * inch, footer_text)
        canvas.restoreState()

    pdf.build(elements, onFirstPage=draw_footer, onLaterPages=draw_footer)

    logger.info(
        "PDF document generated (AI-composed)",
        extra={
            "event": "document_generated",
            "format": "pdf",
            "policy_id": policy_id,
            "ai_composed": True,
        },
    )
    return filepath


# ═══════════════════════════════════════════════════════════════════
#  JSON Export (no AI needed — raw structure)
# ═══════════════════════════════════════════════════════════════════

async def generate_json_export(
    db: AsyncSession, policy_id: str, filename_prefix: str
) -> str:
    """Export the policy structure as formatted JSON. Requires approval."""
    await _check_policy_approved(db, policy_id)

    structure = await _get_latest_structure(policy_id)
    filepath = os.path.join(OUTPUT_DIR, f"{filename_prefix}_{uuid.uuid4().hex[:8]}.json")

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=2, default=str)

    logger.info(
        "JSON document exported",
        extra={"event": "document_generated", "format": "json", "policy_id": policy_id},
    )
    return filepath
